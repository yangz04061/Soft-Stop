
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
基于Word文档生成补充结构的MD模板
提取Word原始排版结构（标题、列表、段落等），生成可直接编辑的模板
自动从文档名称提取标题，无需硬编码
"""

import os
import re
import sys
import glob
from docx import Document

def extract_title_from_filename(filename):
    """
    从Word文件名中提取标题
    例如: "SGMW_VMC HDC_陡坡缓降功能规范 Master_20260302.docx" 
    -> "陡坡缓降功能规范"
    """
    # 移除扩展名
    name = os.path.splitext(filename)[0]
    
    # 移除文件名中的日期（格式：_yyyymmdd 或 _v1等）
    name = re.sub(r'_\d{8}.*$', '', name)  # 移除末尾的日期
    name = re.sub(r'_v\d+.*$', '', name)  # 移除末尾的版本号
    name = re.sub(r'\s+Master.*$', '', name)  # 移除 Master 及之后的内容
    
    # 移除开头的项目代码（SGMW_VMC 等）
    # 保留下划线之后或空格之后的内容
    parts = re.split(r'[_\s]+', name)
    
    # 找到最长的、包含中文或有意义的部分
    candidates = []
    for part in parts:
        # 过滤掉纯英文代码（如HDC、VMC等）且长度小于3的部分
        if len(part) >= 2 and any('\u4e00' <= c <= '\u9fff' for c in part):
            candidates.append(part)
        elif len(part) >= 3 and part.replace(' ', '').replace('-', '').isalpha():
            if not all(c.isupper() for c in part):  # 排除全大写的缩写
                candidates.append(part)
    
    # 如果有候选项，返回最后一个（通常是最重要的描述）；否则返回原始名称的主体部分
    if candidates:
        return ' '.join(candidates)
    
    # 如果没有找到，尝试从最后一个下划线之后提取
    if '_' in name:
        return name.split('_')[-1].strip()
    
    return name.strip()

def get_output_filename(input_file, suffix="_Template"):
    """根据输入文件名生成输出文件名"""
    base_name = os.path.splitext(os.path.basename(input_file))[0]
    # 清理文件名，移除日期和版本号
    base_name = re.sub(r'_\d{8}.*$', '', base_name)
    base_name = re.sub(r'_v\d+.*$', '', base_name)
    base_name = re.sub(r'\s+Master.*$', '', base_name)
    return f"{base_name}{suffix}.md"

def should_include_paragraph(para):
    """判断段落是否应该包含在模板中"""
    text = para.text.strip()
    style_name = para.style.name if para.style else "Normal"
    
    # 包括所有标题、列表项
    if any(x in style_name for x in ["Heading", "List"]):
        return True
    
    # 排除过短的文本
    if len(text) < 3:
        return False
    
    # 排除TOC页（包含制表符）
    if '\t' in text:
        return False
    
    # 包括有意义的段落
    return True


def is_numbered_list(para):
    """检查段落是否为编号列表"""
    pPr = para._element.get_or_add_pPr()
    numPr = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
    return numPr is not None

def get_list_level(para):
    """获取列表项的Word编号级别（0, 1, 2, ...）"""
    pPr = para._element.get_or_add_pPr()
    numPr = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
    
    if numPr is None:
        return None
    
    ilvl_elem = numPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl')
    if ilvl_elem is not None:
        return int(ilvl_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'))
    
    return 0

def format_paragraph_to_md(para, next_para=None, numbered_counter=None):
    """将段落格式化为MD - 考虑级别和编号列表"""
    text = para.text.strip()
    style_name = para.style.name if para.style else "Normal"
    is_bold = any(run.bold for run in para.runs)
    
    if "Heading 1" in style_name:
        return f"## {text}"
    elif "Heading 2" in style_name:
        return f"### {text}"
    elif "Heading 3" in style_name:
        return f"#### {text}"
    elif "List" in style_name:
        # 检查是否为编号列表及其级别
        is_numbered = is_numbered_list(para)
        list_level = get_list_level(para) if is_numbered else None
        
        if is_numbered:
            # 编号列表 - 标记为特殊格式，由调用方处理
            return f"__NUMBERED__{text}"
        else:
            # 无编号的List Paragraph - 作为前一个列表项的延续处理（返回None让后续逻辑处理）
            return None
    else:
        # 正文段落 - 在模板中作为参考
        return None

def find_section_boundaries(doc):
    """查找三个章节的范围"""
    boundaries = {
        'section_2': (None, None),
        'section_3': (None, None),
        'section_5': (None, None),
    }
    
    section_2_start = None
    section_3_start = None
    section_5_start = None
    section_6_start = None
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style_name = para.style.name if para.style else "Normal"
        
        if section_2_start is None and "功能描述" in text and "Heading" in style_name:
            section_2_start = i
        elif section_3_start is None and "系统架构" in text and "Heading 1" in style_name:
            section_3_start = i
        elif section_5_start is None and "功能逻辑策略" in text and "Heading 1" in style_name and i > (section_3_start or 0) + 10:
            section_5_start = i
        elif section_5_start is not None and section_6_start is None and "Heading 1" in style_name and i > section_5_start + 10:
            if "功能降级" in text or "技术指标" in text:
                section_6_start = i
                break
    
    if section_2_start and section_3_start:
        boundaries['section_2'] = (section_2_start, section_3_start)
    if section_3_start and section_5_start:
        boundaries['section_3'] = (section_3_start, section_5_start)
    if section_5_start:
        section_5_end = section_6_start if section_6_start else len(doc.paragraphs)
        boundaries['section_5'] = (section_5_start, section_5_end)
    
    return boundaries

def generate_formatted_template(input_file, output_file=None, doc_title=None):
    """生成格式化模板"""
    doc = Document(input_file)
    boundaries = find_section_boundaries(doc)
    
    # 如果没有提供标题，从文件名提取
    if doc_title is None:
        filename = os.path.basename(input_file)
        doc_title = extract_title_from_filename(filename)
    
    # 如果没有提供输出文件名，自动生成
    if output_file is None:
        output_file = get_output_filename(input_file)
    
    print("=== 文档信息 ===")
    print(f"输入文件: {os.path.basename(input_file)}")
    print(f"提取标题: {doc_title}")
    print(f"输出文件: {output_file}")
    print(f"\n=== 章节范围 ===")
    print(f"Section 2: {boundaries['section_2']}")
    print(f"Section 3: {boundaries['section_3']}")
    print(f"Section 5: {boundaries['section_5']}")
    print()
    
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(f"# {doc_title}\n\n")
        f.write("> **使用说明**: 保留所有标题（##-####）和列表项（- ），将对应内容直接替换为实际内容\n\n")
        
        # Section 2 - 处理多层级编号列表
        if boundaries['section_2'][0] is not None:
            f.write("## 2 功能描述\n\n")
            start, end = boundaries['section_2']
            last_was_heading = False
            numbered_counter = {0: 0, 1: 0, 2: 0, 3: 0}  # 跟踪每个级别的编号
            last_list_level = None  # 记录上一个列表项的级别
            for i in range(start + 1, end):
                para = doc.paragraphs[i]
                if should_include_paragraph(para):
                    next_para = doc.paragraphs[i + 1] if i + 1 < end else None
                    formatted = format_paragraph_to_md(para, next_para)
                    if formatted:
                        # 处理编号列表标记
                        if formatted.startswith("__NUMBERED__"):
                            text = formatted.replace("__NUMBERED__", "")
                            list_level = get_list_level(para)
                            
                            if list_level is not None:
                                # 复位更深层级的计数器
                                for level in range(list_level + 1, 4):
                                    numbered_counter[level] = 0
                                
                                # 增加当前级别计数器
                                numbered_counter[list_level] += 1
                                last_list_level = list_level  # 记录当前列表级别
                                
                                # 根据级别生成相应的MD格式
                                if list_level == 0:
                                    formatted = f"{numbered_counter[0]}. {text}"
                                elif list_level == 1:
                                    formatted = f"   {numbered_counter[1]}. {text}"
                                elif list_level >= 2:
                                    indent_str = "      " + ("   " * (list_level - 2))
                                    formatted = f"{indent_str}- {text}"
                        
                        f.write(formatted + "\n")
                        last_was_heading = formatted.startswith("#") or formatted.startswith("-") or (formatted[0:2].replace(" ", "").isdigit() and formatted[2] == ".")
                    else:
                        text = para.text.strip()
                        style_name = para.style.name if para.style else "Normal"
                        
                        # 如果是Normal段落且前有列表项，作为列表项的延续内容（添加缩进）
                        if text and "Normal" in style_name and last_list_level is not None:
                            # 使用与最后一个列表项相同的缩进
                            if last_list_level >= 2:
                                indent_str = "      " + ("   " * (last_list_level - 2))
                            else:
                                indent_str = "  " * (last_list_level + 1)
                            f.write(f"{indent_str}{text}\n\n")
                        elif text and not last_was_heading:
                            f.write("\n")
                            f.write(text + "\n\n")
                            last_list_level = None  # 重置，因为遇到非列表项内容
                        last_was_heading = False
            f.write("\n")
        
        # Section 3 - 处理多层级编号列表
        if boundaries['section_3'][0] is not None:
            f.write("## 3 系统架构\n\n")
            start, end = boundaries['section_3']
            last_was_heading = False
            numbered_counter = {0: 0, 1: 0, 2: 0, 3: 0}  # 跟踪每个级别的编号
            last_list_level = None  # 记录上一个列表项的级别
            
            for i in range(start + 1, end):
                para = doc.paragraphs[i]
                if should_include_paragraph(para):
                    next_para = doc.paragraphs[i + 1] if i + 1 < end else None
                    formatted = format_paragraph_to_md(para, next_para)
                    
                    if formatted:
                        # 处理编号列表标记
                        if formatted.startswith("__NUMBERED__"):
                            text = formatted.replace("__NUMBERED__", "")
                            list_level = get_list_level(para)
                            
                            if list_level is not None:
                                # 复位更深层级的计数器
                                for level in range(list_level + 1, 4):
                                    numbered_counter[level] = 0
                                
                                # 增加当前级别计数器
                                numbered_counter[list_level] += 1
                                last_list_level = list_level  # 记录当前列表级别
                                
                                # 根据级别生成相应的MD格式
                                if list_level == 0:
                                    # 顶级项（加粗主题）
                                    formatted = f"{numbered_counter[0]}. {text}"
                                elif list_level == 1:
                                    # 第二级（子标题）
                                    formatted = f"   {numbered_counter[1]}. {text}"
                                elif list_level >= 2:
                                    # 第三级及以上（用bullet）
                                    indent_str = "      " + ("   " * (list_level - 2))
                                    formatted = f"{indent_str}- {text}"
                        
                        f.write(formatted + "\n")
                        last_was_heading = formatted.startswith("#") or formatted.startswith("-") or (formatted[0:2].replace(" ", "").isdigit() and formatted[2] == ".")
                    else:
                        text = para.text.strip()
                        style_name = para.style.name if para.style else "Normal"
                        
                        # 如果是Normal段落且前有列表项，作为列表项的延续内容（添加缩进）
                        if text and "Normal" in style_name and last_list_level is not None:
                            # 使用与最后一个列表项相同的缩进
                            if last_list_level >= 2:
                                indent_str = "      " + ("   " * (last_list_level - 2))
                            else:
                                indent_str = "  " * (last_list_level + 1)
                            f.write(f"{indent_str}{text}\n\n")
                        elif text and not last_was_heading:
                            f.write("\n")
                            f.write(text + "\n\n")
                            last_list_level = None  # 重置，因为遇到非列表项内容
                        last_was_heading = False
            f.write("\n")
        
        # Section 5 - 处理多层级编号列表
        if boundaries['section_5'][0] is not None:
            f.write("## 5 功能逻辑策略\n\n")
            start, end = boundaries['section_5']
            last_was_heading = False
            numbered_counter = {0: 0, 1: 0, 2: 0, 3: 0}  # 跟踪每个级别的编号
            last_list_level = None  # 记录上一个列表项的级别
            for i in range(start + 1, end):
                para = doc.paragraphs[i]
                if should_include_paragraph(para):
                    next_para = doc.paragraphs[i + 1] if i + 1 < end else None
                    formatted = format_paragraph_to_md(para, next_para)
                    if formatted:
                        # 处理编号列表标记
                        if formatted.startswith("__NUMBERED__"):
                            text = formatted.replace("__NUMBERED__", "")
                            list_level = get_list_level(para)
                            
                            if list_level is not None:
                                # 复位更深层级的计数器
                                for level in range(list_level + 1, 4):
                                    numbered_counter[level] = 0
                                
                                # 增加当前级别计数器
                                numbered_counter[list_level] += 1
                                last_list_level = list_level  # 记录当前列表级别
                                
                                # 根据级别生成相应的MD格式
                                if list_level == 0:
                                    formatted = f"{numbered_counter[0]}. {text}"
                                elif list_level == 1:
                                    formatted = f"   {numbered_counter[1]}. {text}"
                                elif list_level >= 2:
                                    indent_str = "      " + ("   " * (list_level - 2))
                                    formatted = f"{indent_str}- {text}"
                        
                        f.write(formatted + "\n")
                        last_was_heading = formatted.startswith("#") or formatted.startswith("-") or (formatted[0:2].replace(" ", "").isdigit() and formatted[2] == ".")
                    else:
                        text = para.text.strip()
                        style_name = para.style.name if para.style else "Normal"
                        
                        # 检查是否是"新条件"段落（包含"条件"关键字）- 这种应该重置编号
                        is_new_section = "条件" in text and "：" in text
                        
                        # 如果是Normal段落且前有列表项
                        if text and "Normal" in style_name and last_list_level is not None and not is_new_section:
                            # 不是新条件的Normal → 作为列表项的延续处理（添加缩进）
                            if last_list_level >= 2:
                                indent_str = "      " + ("   " * (last_list_level - 2))
                            else:
                                indent_str = "  " * (last_list_level + 1)
                            f.write(f"{indent_str}{text}\n\n")
                        elif text and not last_was_heading:
                            # 新Normal段落（可能是条件说明或其他）
                            f.write("\n")
                            f.write(text + "\n\n")
                            last_list_level = None  # 重置，因为遇到非列表项内容
                            # 重置编号计数器 - 新的条件/段落需要重新开始编号
                            numbered_counter = {0: 0, 1: 0, 2: 0, 3: 0}
                        last_was_heading = False
    
    print(f"✓ 格式化模板生成完成: {output_file}")

if __name__ == "__main__":
    # 支持从命令行传入参数或交互式输入
    if len(sys.argv) > 1:
        input_file = sys.argv[1]
        output_file = sys.argv[2] if len(sys.argv) > 2 else None
        doc_title = sys.argv[3] if len(sys.argv) > 3 else None
    else:
        # 交互式输入
        print("=== MD 模板生成工具 ===\n")
        
        input_file = input("请输入 Word 文档路径 (或直接按Enter使用当前目录的docx文件): ").strip()
        
        # 如果未指定，查找当前目录的第一个docx文件
        if not input_file:
            docx_files = glob.glob("*.docx")
            if docx_files:
                input_file = docx_files[0]
                print(f"检测到文件: {input_file}")
            else:
                print("错误: 未找到 .docx 文件")
                sys.exit(1)
        
        # 验证文件存在
        if not os.path.exists(input_file):
            print(f"错误: 文件不存在 - {input_file}")
            sys.exit(1)
        
        output_file = input("输出文件名 (直接按Enter自动生成): ").strip() or None
        doc_title = input("文档标题 (直接按Enter从文件名提取): ").strip() or None
    
    try:
        generate_formatted_template(input_file, output_file, doc_title)
        print("\n✓ 完成!")
    except Exception as e:
        print(f"\n✗ 出错: {e}")
        sys.exit(1)
